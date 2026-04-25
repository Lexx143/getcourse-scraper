import re
import os
import sys
import time
import uuid
import argparse
import logging
import subprocess
import http.cookiejar
from urllib.parse import urljoin
from pathlib import Path

# Global state for GUI control
ABORT_DOWNLOAD = False
PAUSE_DOWNLOAD = False
CURRENT_YTDLP_PID = None
CREATED_FILES = []

def check_pause_abort():
    global ABORT_DOWNLOAD, PAUSE_DOWNLOAD
    if ABORT_DOWNLOAD:
        raise InterruptedError("Загрузка отменена пользователем.")
    while PAUSE_DOWNLOAD:
        if ABORT_DOWNLOAD:
            raise InterruptedError("Загрузка отменена пользователем.")
        time.sleep(0.5)

def cleanup_session():
    """Deletes all files created during the current run if aborted."""
    for f in CREATED_FILES:
        try:
            if f.exists():
                f.unlink()
        except:
            pass
    CREATED_FILES.clear()

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

# Optional imports for scraping, so Git init can run even without them
try:
    import requests
    from bs4 import BeautifulSoup, NavigableString
    import docx
    from docx.shared import Inches
except ImportError:
    requests = None
    BeautifulSoup = None
    docx = None
    NavigableString = None

def init_git_repo(base_dir: Path):
    """Initializes Git repository if it doesn't exist, creates boilerplate files and first commit."""
    if (base_dir / '.git').exists():
        return

    logging.info("Initializing new Git repository...")
    
    # Create boilerplate files
    gitignore_path = base_dir / '.gitignore'
    if not gitignore_path.exists():
        gitignore_path.write_text("data/\n__pycache__/\n.env\nvenv/\n*.pyc\ncookies.txt\n", encoding='utf-8')

    readme_path = base_dir / 'README.md'
    if not readme_path.exists():
        readme_path.write_text("# GetCourse Scraper\n\nSee instructions.\n", encoding='utf-8')

    requirements_path = base_dir / 'requirements.txt'
    if not requirements_path.exists():
        requirements_path.write_text("requests\nbeautifulsoup4\npython-docx\nyt-dlp\n", encoding='utf-8')

    data_dir = base_dir / 'data'
    data_dir.mkdir(exist_ok=True)

    try:
        subprocess.run(["git", "init"], cwd=base_dir, check=True, capture_output=True)
        subprocess.run(["git", "add", ".gitignore", "README.md", "requirements.txt", Path(__file__).name], cwd=base_dir, check=True, capture_output=True)
        status = subprocess.run(["git", "status", "--porcelain"], cwd=base_dir, capture_output=True, text=True)
        if status.stdout.strip():
            subprocess.run(["git", "commit", "-m", "Initial commit: Scraper engine structure"], cwd=base_dir, check=True, capture_output=True)
            logging.info("Committed initial repository structure.")
    except subprocess.CalledProcessError as e:
        logging.error(f"Git command failed: {e.stderr.decode('utf-8') if e.stderr else str(e)}")
    except FileNotFoundError:
        logging.error("Git is not installed or not found in PATH.")

def clean_filename(name: str) -> str:
    """Cleans a string to be used as a safe filename."""
    # Убираем слово 'просмотрено', которое GetCourse добавляет к названиям уроков
    name = name.replace('просмотрено', '').strip()
    safe = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_', '(', ')')).strip()
    return safe.replace(' ', '_').strip('_')

def get_session(base_dir: Path):
    """Creates a requests session and loads cookies if available."""
    session = requests.Session()
    headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'}
    session.headers.update(headers)
    
    cookie_file = base_dir / 'cookies.txt'
    if cookie_file.exists():
        try:
            cj = http.cookiejar.MozillaCookieJar(cookie_file)
            cj.load(ignore_discard=True, ignore_expires=True)
            session.cookies.update(cj)
            logging.info("Loaded cookies from cookies.txt")
        except Exception as e:
            logging.error(f"Failed to load cookies.txt: {e}")
    else:
        logging.warning("No cookies.txt found. The scraper might not have access to closed trainings.")
        
    return session

def process_element(elem, doc, session, course_dir):
    """Recursively processes elements to preserve order of text and images, and avoids adding duplicate headings."""
    check_pause_abort()
    
    if isinstance(elem, NavigableString):
        text = str(elem).strip()
        if text:
            if text.startswith('{"signature":') or text.startswith('{"object_type_id":'):
                return
            doc.add_paragraph(text)
        return

    if elem.name in ['script', 'style', 'iframe', 'video']:
        return

    # Filter out user avatars and comment blocks
    if hasattr(elem, 'get'):
        classes = elem.get('class', [])
        if isinstance(classes, list):
            bad_classes = {'user-image', 'user-profile-image', 'answer-comment', 'feedback-modal', 'user-avatar', 'comment-form', 'comments-block', 'lt-form', 'form-group', 'lesson-answer-comment', 'user-related-data', 'vue-component', 'task-block'}
            if any(c in bad_classes for c in classes):
                return
        elif isinstance(classes, str):
            if any(c in classes for c in ['user-image', 'user-profile-image', 'answer-comment', 'feedback-modal', 'user-avatar', 'comment-form', 'comments-block', 'lt-form', 'form-group', 'lesson-answer-comment', 'user-related-data', 'vue-component', 'task-block']):
                return

    # Handle file downloads from <a> tags (Word, PDF, etc.)
    if elem.name == 'a':
        href = elem.get('href')
        if href and ('/fileservice/file/download' in href or href.endswith(('.pdf', '.doc', '.docx', '.xls', '.xlsx', '.zip', '.rar'))):
            link_text = elem.get_text(strip=True)
            if not link_text:
                link_text = "Приложение"
            
            ext = ""
            if "." in href.split("/")[-1]:
                ext = "." + href.split("/")[-1].split(".")[-1].split("?")[0]
            
            if href.startswith('//'):
                href = 'https:' + href
            elif href.startswith('/'):
                href = 'https://niifittech.ru' + href
                
            try:
                logging.info(f"Скачивание прикрепленного файла: {link_text}")
                r = session.get(href, stream=True, timeout=15)
                r.raise_for_status()
                
                cd = r.headers.get('Content-Disposition', '')
                if not ext and 'filename=' in cd:
                    match = re.search(r'filename="?([^";]+)"?', cd)
                    if match:
                        server_filename = match.group(1)
                        if '.' in server_filename:
                            ext = '.' + server_filename.split('.')[-1]
                
                filename = clean_filename(link_text) + ext
                file_path = course_dir / filename
                
                with open(file_path, 'wb') as f:
                    for chunk in r.iter_content(8192):
                        check_pause_abort()
                        f.write(chunk)
                CREATED_FILES.append(file_path)
                
                doc.add_paragraph(f"[Скачан файл: {file_path.name}]", style='Intense Quote')
            except Exception as e:
                logging.error(f"Failed to download attached file {href}: {e}")
                
            # Do not recurse into children of this <a> to skip the PDF/DOCX icon image
            return

    if elem.name == 'img':
        src = elem.get('src')
        if src:
            if src.startswith('//'):
                src = 'https:' + src
            elif src.startswith('/'):
                src = 'https://niifittech.ru' + src
            
            try:
                r = session.get(src, stream=True, timeout=15)
                r.raise_for_status()
                
                temp_img = course_dir / f"temp_{uuid.uuid4().hex}.jpg"
                with open(temp_img, 'wb') as f:
                    for chunk in r.iter_content(1024):
                        f.write(chunk)
                
                try:
                    doc.add_picture(str(temp_img), width=Inches(6))
                except Exception as pic_err:
                    logging.error(f"Failed to insert image to doc: {pic_err}")
                
                if temp_img.exists():
                    os.remove(temp_img)
            except Exception as e:
                logging.error(f"Failed to download image {src}: {e}")
        return

    if elem.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        text = elem.get_text(strip=True)
        if text:
            doc.add_heading(text, level=int(elem.name[1]))
        return

    if elem.name in ['p', 'li', 'blockquote']:
        text = elem.get_text(separator=' ', strip=True)
        if text:
            if elem.name == 'li':
                doc.add_paragraph(text, style='List Bullet')
            else:
                doc.add_paragraph(text)
        
        # Process images inside this block element so they don't get lost
        for img in elem.find_all('img'):
            process_element(img, doc, session, course_dir)
        return

    # Process children recursively for layout preservation
    if hasattr(elem, 'children'):
        for child in elem.children:
            process_element(child, doc, session, course_dir)

def process_block(block, doc, session, course_dir):
    """Processes a single content block, extracting text, images."""
    for child in block.children:
        process_element(child, doc, session, course_dir)

def get_master_playlist(url: str, session) -> str:
    """Extracts masterPlaylistUrl from GetCourse player proxy page to bypass yt-dlp limitations."""
    try:
        r = session.get(url, timeout=15)
        match = re.search(r'"masterPlaylistUrl":"([^"]+)"', r.text)
        if match:
            # Unescape JSON slashes
            return match.group(1).replace('\\/', '/')
    except Exception as e:
        logging.error(f"Failed to extract playlist from {url}: {e}")
    return url

def process_lesson(url: str, session, course_dir: Path, title: str):
    """Fetches a lesson page and extracts its content blocks and videos."""
    clean_title = clean_filename(title)
    logging.info(f"Fetching lesson: {clean_title} ({url})")
    try:
        response = session.get(url, timeout=15)
        response.raise_for_status()
    except requests.RequestException as e:
        logging.error(f"Failed to fetch lesson {url}: {e}")
        return

    soup = BeautifulSoup(response.text, 'html.parser')
    blocks = soup.select('.lite-block-live-wrapper, .v-spacing') 
    
    if not blocks:
        logging.warning(f"No content blocks found in lesson: {clean_title}")
        return

    # Create one document for the whole lesson named after the lesson
    doc_path = course_dir / f"{clean_title}.docx"
    doc = docx.Document()
    doc.add_heading(title.replace('просмотрено', '').strip(), level=1)

    for idx, block in enumerate(blocks, start=1):
        process_block(block, doc, session, course_dir)
        
    doc.save(doc_path)
    CREATED_FILES.append(doc_path)
    logging.info(f"Saved text and images -> {doc_path.name}")

    # Extract videos from ALL blocks
    videos = soup.find_all(lambda tag: tag.name == 'iframe' or tag.has_attr('data-video-hash'))
    video_items = []
    
    for v in videos:
        if v.name == 'iframe' and v.get('src'):
            if 'youtube.com' in v['src'] or 'vimeo.com' in v['src'] or 'player' in v['src']:
                video_items.append(v['src'])
        elif v.has_attr('data-video-hash'):
            video_items.append(f"https://player.getcourse.ru/video/hash/{v['data-video-hash']}")

    # Remove duplicates
    video_items = list(dict.fromkeys(video_items))

    if video_items:
        for idx, item in enumerate(video_items, start=1):
            target = item
            
            # If it's a GetCourse player proxy, fetch the actual m3u8 playlist to avoid yt-dlp extractor failures
            if 'sign-player' in target or 'player.getcourse.ru' in target:
                playlist_url = get_master_playlist(target, session)
                if playlist_url:
                    target = playlist_url

            logging.info(f"Downloading video {idx}/{len(video_items)} for {clean_title}...")
            
            # Use lesson title in the video filename
            video_name = f"{clean_title}_{idx}" if len(video_items) > 1 else clean_title
            out_tmpl = str(course_dir / f"{video_name}.%(ext)s")
            
            # Predict the final video file path assuming mp4
            CREATED_FILES.append(course_dir / f"{video_name}.mp4")
            
            global CURRENT_YTDLP_PID
            cmd = ["yt-dlp", "--cookies", "cookies.txt", "--referer", "https://niifittech.ru/", "-o", out_tmpl, target]
            
            try:
                proc = subprocess.Popen(cmd, cwd=Path(__file__).parent, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True)
                CURRENT_YTDLP_PID = proc.pid
                
                for line in proc.stdout:
                    check_pause_abort()
                    # Pass the yt-dlp output to logger with a special prefix
                    logging.info(f"YTDLP_OUT: {line.strip()}")
                
                proc.wait()
                CURRENT_YTDLP_PID = None
                
                if proc.returncode == 0:
                    logging.info(f"Successfully downloaded video for: {clean_title}")
                else:
                    logging.error(f"Failed to download video. yt-dlp returned error. Moving on...")
            except InterruptedError as e:
                if CURRENT_YTDLP_PID:
                    try:
                        import psutil
                        psutil.Process(CURRENT_YTDLP_PID).kill()
                    except:
                        pass
                CURRENT_YTDLP_PID = None
                raise e
            except FileNotFoundError:
                CURRENT_YTDLP_PID = None
                logging.error("yt-dlp is not installed! Cannot download video. Please install it.")
        
    check_pause_abort()
    time.sleep(1) # Polite delay between lessons

def crawl_course(url: str, session, data_dir: Path, visited=None):
    """Crawls a stream/course page, finds lessons and substreams, and processes them recursively."""
    if visited is None:
        visited = set()
    
    # Normalize URL to avoid infinite loops with trailing slashes
    norm_url = url.split('?')[0].rstrip('/')
    if norm_url in visited:
        return
    visited.add(norm_url)

    logging.info(f"Crawling course stream: {url}")
    try:
        check_pause_abort()
        response = session.get(url, timeout=15)
        response.raise_for_status()
    except requests.RequestException as e:
        logging.error(f"Failed to fetch course {url}: {e}")
        return

    soup = BeautifulSoup(response.text, 'html.parser')
    
    title_elem = soup.find('h1')
    course_title = title_elem.get_text(strip=True) if title_elem else "Course"
    
    # This acts as the module/theme directory
    course_dir = data_dir / clean_filename(course_title)
    course_dir.mkdir(parents=True, exist_ok=True)
    
    logging.info(f"Directory created/used: {course_dir.name}")
    
    # Links to sub-streams/modules
    substream_links = soup.find_all('a', href=lambda href: href and '/teach/control/stream/view/id/' in href)
    
    # Links to actual lessons
    lesson_links = soup.find_all('a', href=lambda href: href and ('/lesson/view/id/' in href or '/teach/control/lesson/view/id/' in href))
    
    # 1. Process substreams recursively
    for a in substream_links:
        href = a['href']
        full_url = urljoin(url, href)
        norm_full = full_url.split('?')[0].rstrip('/')
        if norm_full not in visited and norm_full != norm_url:
            crawl_course(full_url, session, course_dir, visited)

    # 2. Process lessons
    seen_lessons = set()
    lessons = []
    for a in lesson_links:
        href = a['href']
        norm_href = href.split('?')[0]
        if norm_href not in seen_lessons:
            seen_lessons.add(norm_href)
            title = a.get_text(strip=True)
            if not title:
                parent = a.find_parent('div', class_='stream-title') or a.find_parent('td')
                if parent:
                    title = parent.get_text(strip=True)
            if not title:
                title = f"Lesson_{len(lessons)+1}"
                
            full_url = urljoin(url, href)
            lessons.append((full_url, title))

    if not lesson_links and not substream_links:
        logging.warning("No lesson or substream links found. Is this a valid page? Ensure cookies are valid.")
        # Only fallback if this is NOT the root call of the recursive crawl
        if len(visited) == 1:
            process_lesson(url, session, data_dir, "Single_Lesson")
        return

    if lessons:
        logging.info(f"Found {len(lessons)} lessons in this stream.")
        for idx, (lesson_url, lesson_title) in enumerate(lessons, start=1):
            logging.info(f"[{idx}/{len(lessons)}] Processing {lesson_title}...")
            numbered_title = f"{idx:03d}_{lesson_title}"
            process_lesson(lesson_url, session, course_dir, numbered_title)

def scrape(url_or_path: str):
    """Main scraping orchestrator."""
    if requests is None or BeautifulSoup is None or docx is None:
        logging.error("Missing dependencies. Please run: pip install -r requirements.txt")
        sys.exit(1)

    base_dir = Path(__file__).parent.resolve()
    data_dir = base_dir / "data"
    data_dir.mkdir(exist_ok=True)
    
    session = get_session(base_dir)
    
    if url_or_path.startswith("http://") or url_or_path.startswith("https://"):
        crawl_course(url_or_path, session, data_dir)
    else:
        path = Path(url_or_path)
        if not path.exists():
            logging.error(f"File not found: {path}")
            return
        logging.info(f"Reading local file: {path}")
        html_content = path.read_text(encoding='utf-8')
        soup = BeautifulSoup(html_content, 'html.parser')
        blocks = soup.select('.lite-block-live-wrapper, .v-spacing')
        
        # Local file mockup
        doc_path = data_dir / "Local_File.docx"
        doc = docx.Document()
        doc.add_heading("Local_File", level=1)
        for idx, block in enumerate(blocks, start=1):
            process_block(block, doc, session, data_dir)
        doc.save(doc_path)

def main():
    parser = argparse.ArgumentParser(description="GetCourse clean and scalable scraper engine.")
    parser.add_argument('--url', type=str, help="URL of the stream/course to scrape")
    args = parser.parse_args()

    base_dir = Path(__file__).parent.resolve()
    
    init_git_repo(base_dir)

    if args.url:
        scrape(args.url)
    else:
        logging.info("No URL provided. Git structure initialized. Use --url <course_stream_url> to start scraping.")

if __name__ == "__main__":
    main()
